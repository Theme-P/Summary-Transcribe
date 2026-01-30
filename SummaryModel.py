import requests
import json 
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# NTC AI Gateway API configuration
NTC_API_KEY = os.getenv("NTC_API_KEY")
NTC_API_URL = os.getenv("NTC_API_URL", "https://aigateway.ntictsolution.com/v1/chat/completions")


def summarize_transcription(transcription_text: str, language: str = "Thai") -> str:
    """
    Summarize transcription text from WhisperX using GPT-4o via NTC AI Gateway.
    
    Args:
        transcription_text: The transcription text to summarize
        language: The language for the summary output (default: Thai)
    
    Returns:
        Summarized text
    """
    if not NTC_API_KEY:
        return "Error: NTC_API_KEY not found in environment variables"
    
    headers = {
        "Authorization": f"Bearer {NTC_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "system",
                "content": f"""คุณคือผู้เชี่ยวชาญสรุปการประชุม ทำตามขั้นตอน:
1. วิเคราะห์ประเภทการประชุม
2. สรุปตามโครงสร้างที่เหมาะสม

**ประเภทการประชุม:**
| ประเภท | โครงสร้างหลัก |
|--------|--------------|
| Shareholder Meeting | วาระ → มติ → เงินปันผล → ข้อสรุป |
| Board Meeting | นโยบาย → การอนุมัติ → มติคณะกรรมการ |
| Planning Meeting | เป้าหมาย → แผนงาน → ไทม์ไลน์ → ผู้รับผิดชอบ → ความเสี่ยง |
| Progress Update | สถานะโครงการ → ความคืบหน้า → ปัญหา → แนวทางแก้ → งานถัดไป |
| Strategy Meeting | ทิศทางธุรกิจ → การวิเคราะห์ → กลยุทธ์ → Action Plan |
| Incident Review | รายละเอียดปัญหา → สาเหตุ → ผลกระทบ → แนวทางแก้ไข → การป้องกัน |
| Client Meeting | ข้อเสนอ → Feedback → ข้อตกลง → Next Steps |
| Workshop | หัวข้อ → เนื้อหาสำคัญ → บทเรียน → Action Items |
| Executive Meeting | ประเด็นสำคัญ → การตัดสินใจ → มติ → ผู้รับผิดชอบ |
| Team Meeting | อัพเดตงาน → การมอบหมาย → ปัญหา → สิ่งที่ต้องทำ |
| General Meeting | วาระ → ประเด็นหารือ → ข้อเสนอแนะ → มติ |

**Output Format:**
**[ประเภท]: [หัวข้อการประชุม]**
(สรุปตามโครงสร้างของประเภทนั้น)

**กฎ:** ใช้ภาษา{language} | ใช้ bullet points | แยกตามทีม/คน | ระบุผู้รับผิดชอบ+กำหนดเวลา | ข้ามหัวข้อที่ไม่มีข้อมูล | สรุปมติท้ายสุด"""
            },
            {
                "role": "user",
                "content": f"สรุปการประชุม:\n\n{transcription_text}"
            }
        ],
        "temperature": 0.4,
        "max_tokens": 4000
    }
    
    try:
        response = requests.post(NTC_API_URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        
        result = response.json()
        return result["choices"][0]["message"]["content"]
        
    except requests.exceptions.RequestException as e:
        return f"Error calling NTC API: {str(e)}"
    except (KeyError, IndexError) as e:
        return f"Error parsing response: {str(e)}"


def summarize_from_whisperx_result(result: dict) -> str:
    """
    Summarize from WhisperX result dictionary.
    
    Args:
        result: WhisperX result dictionary containing 'segments'
    
    Returns:
        Summarized text
    """
    # Extract text from all segments
    if "segments" in result:
        full_text = " ".join([seg.get("text", "") for seg in result["segments"]])
    else:
        full_text = str(result)
    
    return summarize_transcription(full_text)


def summarize_with_diarization(
    transcript_with_speakers: str,
    speaker_summary: dict,
    language: str = "Thai"
) -> str:
    """
    Summarize transcription with speaker diarization data for enhanced analysis.
    
    Args:
        transcript_with_speakers: Full transcript with speaker labels
        speaker_summary: Dict with 'speaking_time' and 'word_count' per speaker
        language: Output language (default: Thai)
    
    Returns:
        Detailed summary with speaker analysis
    """
    if not NTC_API_KEY:
        return "Error: NTC_API_KEY not found in environment variables"
    
    # Build speaker info string
    speakers_time = speaker_summary.get('speaking_time', {})
    speakers_words = speaker_summary.get('word_count', {})
    total_time = sum(speakers_time.values()) if speakers_time else 1
    
    speaker_info_lines = []
    for speaker, time_sec in sorted(speakers_time.items(), key=lambda x: -x[1]):
        pct = (time_sec / total_time * 100) if total_time > 0 else 0
        words = speakers_words.get(speaker, 0)
        mins = int(time_sec // 60)
        secs = int(time_sec % 60)
        speaker_info_lines.append(f"- {speaker}: {mins}:{secs:02d} ({pct:.1f}%), {words} คำ")
    
    speaker_info = "\n".join(speaker_info_lines)
    num_speakers = len(speakers_time)
    
    headers = {
        "Authorization": f"Bearer {NTC_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "system",
                "content": f"""คุณคือผู้เชี่ยวชาญวิเคราะห์และสรุปการประชุม

**ขั้นตอน:**
1. วิเคราะห์ข้อมูลผู้พูดเพื่อระบุบทบาท (ประธาน/ผู้นำเสนอ/ผู้เข้าร่วม)
2. วิเคราะห์เนื้อหาเพื่อระบุประเภทการประชุม
3. สรุปตามโครงสร้างที่เหมาะสม
4. **ระบุการสั่งงาน/คำถาม/ข้อตกลง พร้อมชื่อผู้พูดเสมอ**

**ประเภทการประชุม:**
| ประเภท | โครงสร้าง |
|--------|----------|
| Shareholder Meeting | วาระ → มติ → เงินปันผล |
| Board Meeting | นโยบาย → การอนุมัติ → มติ |
| Planning Meeting | เป้าหมาย → แผนงาน → ไทม์ไลน์ → ผู้รับผิดชอบ |
| Progress Update | สถานะ → ความคืบหน้า → ปัญหา → แนวทางแก้ → งานถัดไป |
| Strategy Meeting | ทิศทาง → กลยุทธ์ → Action Plan |
| Incident Review | ปัญหา → สาเหตุ → แนวทางแก้ → การป้องกัน |
| Client Meeting | ข้อเสนอ → Feedback → ข้อตกลง → Next Steps |
| Workshop | หัวข้อ → เนื้อหา → บทเรียน → Action Items |
| Executive Meeting | การตัดสินใจ → มติ → ผู้รับผิดชอบ |
| Team Meeting | อัพเดต → มอบหมาย → ปัญหา → สิ่งที่ต้องทำ |
| General Meeting | วาระ → หารือ → ข้อเสนอแนะ → มติ |

**Output Format:**
**[ประเภท]: [หัวข้อ]**

**👥 ผู้เข้าร่วมประชุม ({num_speakers} คน):**
(วิเคราะห์บทบาทจากเนื้อหาการพูด)

**📋 สรุปการประชุม:**
(ตามโครงสร้างของประเภทนั้น)

**📌 การสั่งงาน/มอบหมาย:** (ถ้ามี)
- **[ผู้สั่ง]** สั่งให้ **[ผู้รับมอบหมาย]** ทำ: [เนื้อหา] (กำหนด: [วันที่/เวลา ถ้ามี])

**❓ คำถามสำคัญ:** (ถ้ามี)
- **[ผู้ถาม]** ถาม: "[คำถาม]" → **[ผู้ตอบ]**: "[คำตอบ]"

**✅ ข้อตกลง/มติ:** (ถ้ามี)
- [เนื้อหาข้อตกลง] (เสนอโดย: **[ผู้เสนอ]**)

**กฎสำคัญ:**
- ภาษา{language}
- ใช้ bullet points
- **ต้องระบุชื่อผู้พูด (เช่น คนพูด 1, คนพูด 2) ในทุกการสั่งงาน/คำถาม/ข้อตกลง**
- ระบุผู้รับผิดชอบ+กำหนดเวลาเมื่อมีการมอบหมายงาน
- สรุปมติท้ายสุด"""
            },
            {
                "role": "user",
                "content": f"""**ข้อมูลผู้พูด:**
{speaker_info}

**เนื้อหาการประชุม:**
{transcript_with_speakers}"""
            }
        ],
        "temperature": 0.4,
        "max_tokens": 4000
    }
    
    try:
        response = requests.post(NTC_API_URL, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        
        result = response.json()
        return result["choices"][0]["message"]["content"]
        
    except requests.exceptions.RequestException as e:
        return f"Error calling NTC API: {str(e)}"
    except (KeyError, IndexError) as e:
        return f"Error parsing response: {str(e)}"


def export_to_docx(
    summary_text: str,
    output_path: str,
    audio_file: str = None,
    processing_time: dict = None,
    speaker_summary: dict = None
) -> str:
    """
    Export summary to a formatted DOCX file.
    
    Args:
        summary_text: The summary text (markdown format)
        output_path: Path for the output DOCX file
        audio_file: Original audio file name (optional)
        processing_time: Dict with processing times (optional)
        speaker_summary: Dict with speaker info (optional)
    
    Returns:
        Path to the created DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import re
        from datetime import datetime
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"
    
    doc = Document()
    
    # Title
    title = doc.add_heading('📝 สรุปการประชุม', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Main summary content (skip metadata and speaker summary)
    
    # Parse markdown and add to document
    lines = summary_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Handle headers (**, ##, etc.)
        if line.startswith('**') and line.endswith('**'):
            # Bold header
            text = line.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith('##'):
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=2)
        elif line.startswith('#'):
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=1)
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point
            text = line[2:].strip()
            # Handle bold text within bullet points
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
    
    # Save document
    doc.save(output_path)
    return output_path


def _add_formatted_text(paragraph, text: str):
    """Helper function to add text with markdown bold formatting to a paragraph."""
    import re
    
    # Split by bold markers (**text**)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Regular text
            paragraph.add_run(part)

